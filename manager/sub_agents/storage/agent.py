from google.adk.agents import Agent
from manager.db.contract_db import create_contract
from docx import Document
import os
from datetime import datetime

# def save_contract(input: dict):
#     contract_id = input.get("contract_id", f"default_{datetime.utcnow().timestamp()}")
#     contract_text = input.get("contract", "")
#     domain = input.get("domain", "Unknown")

#     # Save to MongoDB using your own create_contract()
#     create_contract({
#         "contract_id": contract_id,
#         "contract": contract_text,
#         "domain": domain,
#         "created_by": input.get("created_by", "unknown")
#     })

#     # Save to Downloads
#     downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
#     file_path = os.path.join(downloads_dir, f"{contract_id}.docx")

#     doc = Document()
#     doc.add_heading(f"{domain} Contract", 0)
#     doc.add_paragraph(contract_text)
#     doc.save(file_path)

#     return f"✅ Contract '{contract_id}' saved successfully."
def safe_filename(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()

def save_contract(input: dict):
    contract_id = input.get("contract_id", f"default_{datetime.utcnow().timestamp()}")
    contract_id = safe_filename(contract_id) or "contract"  # ✅ Fix invalid filename

    contract_text = input.get("contract", "")
    domain = input.get("domain", "Unknown")

    # Save to MongoDB
    create_contract({
        "contract_id": contract_id,
        "contract": contract_text,
        "domain": domain,
        "created_by": input.get("created_by", "unknown")
    })

    # Save to Downloads folder
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    file_path = os.path.join(downloads_dir, f"{contract_id}.docx")

    doc = Document()
    doc.add_heading(f"{domain} Contract", 0)
    doc.add_paragraph(contract_text)
    doc.save(file_path)

    return f"✅ Contract '{contract_id}' saved successfully."



storage_agent = Agent(
    name="storage",
    model="gemini-2.0-flash",
    description="Stores contract to MongoDB and Downloads.",
    instruction="""
You are the storage agent in the CLM system.

You will receive a contract dictionary with:
- contract_id
- contract (text)
- domain
- created_by

Save this contract:
1. Into MongoDB (via backend)
2. As a .docx file inside user's Downloads folder

Respond with: ✅ Contract saved successfully (or error if failed).
""",
   tools=[save_contract],
)


